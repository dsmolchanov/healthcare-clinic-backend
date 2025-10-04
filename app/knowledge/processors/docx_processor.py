"""
DOCX Document Processor using python-docx
"""

import io
import logging
from typing import Dict, Any, List

try:
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False
    logging.warning("python-docx not available. Install with: pip install python-docx")

from ..router import BaseProcessor, InputData, ProcessedDocument

logger = logging.getLogger(__name__)

class DocxProcessor(BaseProcessor):
    """Process DOCX documents and extract text, structure, and formatting"""
    
    supported_mime_types = [
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'application/msword'
    ]
    
    def can_process(self, mime_type: str) -> bool:
        """Check if this processor can handle the given mime type"""
        return mime_type in self.supported_mime_types and PYTHON_DOCX_AVAILABLE
    
    async def process(self, input_data: InputData) -> ProcessedDocument:
        """Process DOCX document and extract structured content"""
        if not PYTHON_DOCX_AVAILABLE:
            return await self._process_fallback(input_data)
        
        try:
            # Load document from bytes
            doc_bytes = input_data.content if isinstance(input_data.content, bytes) else input_data.content.encode()
            doc = Document(io.BytesIO(doc_bytes))
            
            # Extract content
            full_text = ""
            structured_content = []
            metadata = self._extract_metadata(doc)
            tables_data = []
            
            # Process paragraphs
            for para_index, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():
                    para_data = {
                        'index': para_index,
                        'text': paragraph.text,
                        'style': paragraph.style.name if paragraph.style else '',
                        'is_heading': self._is_heading(paragraph),
                        'alignment': str(paragraph.alignment) if paragraph.alignment else '',
                        'runs': []
                    }
                    
                    # Extract run-level formatting
                    for run in paragraph.runs:
                        if run.text:
                            para_data['runs'].append({
                                'text': run.text,
                                'bold': run.bold,
                                'italic': run.italic,
                                'underline': run.underline,
                                'font_name': run.font.name,
                                'font_size': run.font.size.pt if run.font.size else None,
                            })
                    
                    structured_content.append(para_data)
                    full_text += paragraph.text + "\n\n"
            
            # Process tables
            for table_index, table in enumerate(doc.tables):
                table_data = self._extract_table(table, table_index)
                tables_data.append(table_data)
                
                # Add table content to full text
                full_text += f"\n\n[Table {table_index + 1}]\n"
                for row in table_data['rows']:
                    full_text += " | ".join(row) + "\n"
                full_text += "\n"
            
            # Extract lists
            lists = self._extract_lists(doc)
            
            # Create chunks
            chunks = self._create_chunks(structured_content, full_text)
            
            # Extract facts
            facts = self._extract_facts(full_text, metadata, tables_data)
            
            # Count tokens
            tokens_count = len(full_text.split()) * 1.3
            
            return ProcessedDocument(
                content=full_text,
                chunks=chunks,
                facts=facts,
                metadata={
                    **metadata,
                    'paragraph_count': len(doc.paragraphs),
                    'table_count': len(tables_data),
                    'list_count': len(lists),
                    'has_images': self._has_images(doc),
                    'structure_blocks': len(structured_content)
                },
                processing_time_ms=0,
                tokens_count=int(tokens_count)
            )
            
        except Exception as e:
            logger.error(f"Error processing DOCX: {str(e)}")
            raise
    
    def _extract_metadata(self, doc: 'Document') -> Dict[str, Any]:
        """Extract document metadata"""
        metadata = {}
        
        try:
            core_props = doc.core_properties
            metadata = {
                'title': core_props.title or '',
                'author': core_props.author or '',
                'subject': core_props.subject or '',
                'keywords': core_props.keywords or '',
                'created': str(core_props.created) if core_props.created else '',
                'modified': str(core_props.modified) if core_props.modified else '',
                'last_modified_by': core_props.last_modified_by or '',
                'revision': core_props.revision,
                'category': core_props.category or '',
                'comments': core_props.comments or '',
            }
        except Exception as e:
            logger.warning(f"Failed to extract DOCX metadata: {e}")
        
        return metadata
    
    def _is_heading(self, paragraph) -> bool:
        """Check if paragraph is a heading"""
        if paragraph.style:
            return paragraph.style.name.startswith('Heading')
        return False
    
    def _extract_table(self, table, table_index: int) -> Dict[str, Any]:
        """Extract table data"""
        table_data = {
            'index': table_index,
            'rows': [],
            'row_count': len(table.rows),
            'column_count': len(table.columns) if hasattr(table, 'columns') else 0,
        }
        
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                row_data.append(cell_text)
            table_data['rows'].append(row_data)
        
        return table_data
    
    def _extract_lists(self, doc) -> List[Dict[str, Any]]:
        """Extract list items from document"""
        lists = []
        current_list = None
        
        for paragraph in doc.paragraphs:
            # Check if paragraph is a list item (simplified check)
            if paragraph.style and ('List' in paragraph.style.name or 'Bullet' in paragraph.style.name):
                if current_list is None:
                    current_list = {
                        'type': 'bullet' if 'Bullet' in paragraph.style.name else 'numbered',
                        'items': []
                    }
                current_list['items'].append(paragraph.text)
            else:
                if current_list:
                    lists.append(current_list)
                    current_list = None
        
        if current_list:
            lists.append(current_list)
        
        return lists
    
    def _has_images(self, doc) -> bool:
        """Check if document contains images"""
        try:
            # Check for inline shapes (images)
            for paragraph in doc.paragraphs:
                if paragraph._element.xpath('.//w:drawing'):
                    return True
            return False
        except:
            return False
    
    def _create_chunks(self, structured_content: List[Dict], full_text: str) -> List[Dict[str, Any]]:
        """Create intelligent chunks from structured content"""
        chunks = []
        current_chunk = ""
        current_section = ""
        chunk_index = 0
        
        for item in structured_content:
            text = item.get('text', '').strip()
            if not text:
                continue
            
            # Check if it's a heading
            if item.get('is_heading'):
                # Save current chunk if it exists
                if current_chunk:
                    chunks.append({
                        'index': chunk_index,
                        'content': current_chunk.strip(),
                        'section': current_section,
                        'tokens': len(current_chunk.split()),
                    })
                    chunk_index += 1
                    current_chunk = ""
                
                current_section = text
            
            current_chunk += text + "\n\n"
            
            # Check chunk size
            if len(current_chunk.split()) > 500:
                chunks.append({
                    'index': chunk_index,
                    'content': current_chunk.strip(),
                    'section': current_section,
                    'tokens': len(current_chunk.split()),
                })
                chunk_index += 1
                current_chunk = ""
        
        # Add remaining chunk
        if current_chunk:
            chunks.append({
                'index': chunk_index,
                'content': current_chunk.strip(),
                'section': current_section,
                'tokens': len(current_chunk.split()),
            })
        
        return chunks
    
    def _extract_facts(self, text: str, metadata: Dict[str, Any], tables: List[Dict]) -> Dict[str, Any]:
        """Extract structured facts from the document"""
        facts = {}
        
        # Add metadata as facts
        if metadata.get('title'):
            facts['document_title'] = metadata['title']
        if metadata.get('author'):
            facts['document_author'] = metadata['author']
        
        # Extract from tables (often contain structured data)
        if tables:
            facts['tables_summary'] = []
            for table in tables[:3]:  # First 3 tables
                if table['rows']:
                    # Assume first row might be headers
                    facts['tables_summary'].append({
                        'headers': table['rows'][0] if table['rows'] else [],
                        'row_count': table['row_count'],
                        'column_count': table['column_count']
                    })
        
        # Extract common patterns
        import re
        
        # Phone numbers
        phones = re.findall(r'[\+]?[(]?[0-9]{3}[)]?[-\s\.]?[(]?[0-9]{3}[)]?[-\s\.]?[0-9]{4,6}', text)
        if phones:
            facts['phone_numbers'] = list(set(phones))
        
        # Emails
        emails = re.findall(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', text)
        if emails:
            facts['emails'] = list(set(emails))
        
        # Dates (simple pattern)
        dates = re.findall(r'\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b', text)
        if dates:
            facts['dates_mentioned'] = list(set(dates))[:5]  # First 5 unique dates
        
        return facts
    
    async def _process_fallback(self, input_data: InputData) -> ProcessedDocument:
        """Fallback processing when python-docx is not available"""
        text = "DOCX processing requires python-docx. Please install with: pip install python-docx"
        
        if isinstance(input_data.content, str):
            text = input_data.content
        
        return ProcessedDocument(
            content=text,
            chunks=[{
                'index': 0,
                'content': text,
                'section': '',
                'tokens': len(text.split()),
            }],
            facts={},
            metadata={'processor': 'fallback'},
            processing_time_ms=0,
            tokens_count=len(text.split())
        )